from docx import Document
import subprocess
from pathlib import Path

def read_docx(file_path):
    try:
        doc = Document(file_path)
        content = []
        
        # 读取段落文本
        for para in doc.paragraphs:
            content.append(para.text)
        
        # 读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content.append(cell.text)
        
        return "\n".join(content)
    except Exception as e:
        return None


def read_doc_by_antiword(file_path):
    try:
        result = subprocess.run(["antiword", file_path],
                              stdout=subprocess.PIPE,
                              stderr=subprocess.PIPE,
                              text=True)
        return result.stdout
    except Exception as e:
        return None


def read_txt_safe(file_path, max_size_mb=1):
    try:
        path = Path(file_path)
        if not path.exists():
            return None

        size_mb = path.stat().st_size / (1024 ** 2)

        if size_mb > max_size_mb:
            return None

        return path.read_text(encoding='utf-8')

    except Exception as e:
        print(f"读取失败：{str(e)}")
        return None


def main():
    #print(read_docx("test.docx"))
    #print(read_doc_by_antiword("test2.doc"))
    print(read_txt_safe("test.txt"))

if __name__ == '__main__':
    main()
